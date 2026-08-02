import json
import os
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\Java\MinecraftMod\DarkGrey")
RPG_ROOT = Path(r"E:\Java\MinecraftMod\RPGItem")
JSON_PATH = ROOT / "src/main/resources/assets/dark_grey/data/rpg_items.json"
OUTPUT_PATH = RPG_ROOT / "DarkGrey武器道具装备与印记系统说明.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
NAVY = "1F4D78"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "68707A"
DARK = "20262E"
WHITE = "FFFFFF"
GOLD = "B07A16"
RED = "A33A32"
GREEN = "2E7D4A"
PURPLE = "6A4C93"


def read_version() -> str:
    text = (ROOT / "gradle.properties").read_text(encoding="utf-8")
    match = re.search(r"^\s*modVersion\s*=\s*(.+?)\s*$", text, flags=re.MULTILINE)
    return match.group(1).strip() if match else "未知"


def ticks_to_seconds(value):
    return f"{float(value) / 20:g} 秒"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table, color="CAD2DC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, italic=None, font=FONT_CN):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MID_GRAY)


def style_paragraph(paragraph, before=0, after=6, line=1.25, keep_with_next=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next


def add_label_paragraph(doc, label, text, color=NAVY):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=5, line=1.2)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=10.5, bold=True, color=color)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=DARK)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(paragraph, after=4, line=1.25)
    for run in paragraph.runs:
        set_run_font(run, size=10.5, color=DARK)
    if not paragraph.runs:
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5, color=DARK)
    else:
        paragraph.runs[0].text = text
    return paragraph


def add_note_box(doc, title, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=0, line=1.2)
    title_run = paragraph.add_run(title + "  ")
    set_run_font(title_run, size=10.5, bold=True, color=accent)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, size=10.5, color=DARK)
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=0)


def add_meta_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
        set_cell_shading(cells[0], LIGHT_BLUE)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=NAVY)
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=DARK)
        style_paragraph(cells[0].paragraphs[0], after=0, line=1.15, keep_with_next=True)
        style_paragraph(cells[1].paragraphs[0], after=0, line=1.15, keep_with_next=True)
    set_table_geometry(table, [1700, 7660], indent_dxa=120)
    set_table_borders(table)
    return table


def enchantments_text(item):
    raw = item.get("enchantments") or ""
    names = {
        "16": "锋利",
        "21": "抢夺",
        "34": "耐久",
        "51": "无限",
    }
    result = []
    for part in raw.split(","):
        part = part.strip()
        if not part:
            continue
        enchant_id, level = part.split(":", 1)
        result.append(f"{names.get(enchant_id, '附魔' + enchant_id)} {level}")
    return "、".join(result) if result else "无"


def item_details():
    return {
        "bone_flask": {
            "category": "投掷道具",
            "skills": [
                "右键投掷；直接命中造成 12 点伤害，并留下 3×3 骨刺场。",
                "踩中骨刺的敌人受到 2 点伤害并叠加 1 层骨折；骨折稳定期 100 Tick（5 秒）。",
                "骨刺场配置持续 1200 Tick（60 秒）；物品不可堆叠。",
            ],
            "lore": "“掷出后，命中造成伤害并留下骨刺场；踩中骨刺的敌人受伤并叠加骨折。”",
        },
        "bone_crusher": {
            "category": "近战武器 · 剑",
            "skills": [
                "连续对同一目标完成 3 次有效近战命中后，施加 1 层骨折。",
                "每次触发显式使用 100 Tick（5 秒）骨折稳定期，并显示第三击反馈。",
            ],
            "lore": "“粉碎之骨：每有效近战命中 3 次，施加 1 层骨折印记。”",
        },
        "erebus": {
            "category": "法杖",
            "skills": [
                "右键以施法者身体中心为球心，对球体与目标碰撞箱相交的合法生物施加剧毒。",
                "初始半径 3 格，每次使用后增加 1 格，最高 7 格；冷却 20 Tick（1 秒）。",
                "连续 100 Tick（5 秒）不使用时，半径重置为 3 格。",
                "每次必定施加 4 层剧毒；75% 概率额外 +6 层；40% 概率再次额外 +7 层。两次额外判定相互独立。",
                "剧毒稳定期 200 Tick（10 秒），范围不要求隔墙可见。",
            ],
            "lore": "“高举法杖施放深渊诅咒，为球形领域内的目标施加剧毒印记。”",
        },
        "cleave_sword": {
            "category": "近战武器 · 巨剑",
            "skills": [
                "重击：每 5 秒使下一次直接近战命中额外追加“当前手持武器攻击力 ×4”的伤害。",
                "冷却结束后，服务端会向持有者提示重击已经就绪。",
            ],
            "lore": "“重击：蓄势完成后，下一次近战命中爆发出四倍武器攻击力的追加伤害。”",
        },
        "vampire_blade": {
            "category": "近战武器 · 剑",
            "skills": [
                "吸血：将实际造成伤害的 15% 转化为自身生命值。",
                "重击：每 5 秒使下一次直接近战命中额外追加“当前手持武器攻击力 ×4”的伤害。",
            ],
            "lore": "“吸血鬼之刃：以敌人的鲜血延续自身，并在蓄势后斩出沉重一击。”",
        },
        "rainbow_bow": {
            "category": "远程武器 · 弓",
            "skills": [
                "虹之愿采用三段蓄力：1 秒造成 20 点、2 秒造成 30 点、3 秒造成 50 点伤害。",
                "蓄力阶段阈值固定为 20 / 40 / 60 Tick。",
            ],
            "lore": "“虹之愿｜多段蓄力魔法阵。蓄力越完整，箭矢承载的威力越高。”",
        },
        "law_of_cycles": {
            "category": "远程武器 · 弓",
            "skills": [
                "必须滞空蓄力；落地会打断技能。蓄力 100 Tick（5 秒）后圆环完整。",
                "消耗光之矢，释放 30 支水平箭阵；每支光矢造成 150 点基础伤害。",
                "箭阵以视线为中心组成 6×5 阵列，并沿瞄准方向近似平行前进。",
            ],
            "lore": "“圆环之理｜万物起源的永恒循环。必须滞空时才能展开圆环，蓄力至圆环完整。”",
        },
        "arrow_of_light": {
            "category": "弹药",
            "skills": [
                "圆环之理的专用消耗品；自身没有独立技能组件。",
                "圆环之理释放时，用于生成 30 支、每支 150 点伤害的光矢箭阵。",
            ],
            "lore": "当前未配置独立技能 Lore；其定位是“圆环之理”的仪式弹药。",
        },
        "calamity_scythe": {
            "category": "近战武器 · 镰刀",
            "skills": [
                "血祭：每次攻击有 15% 概率造成原基础伤害 3～6 倍的物理暴击。",
                "血祭触发时反噬使用者最大生命值的 5%，生命不足时可能致死。",
                "劫难：右键释放以玩家为圆心、半径 5 格的 360° 横扫；冷却 60 Tick（3 秒）。",
            ],
            "lore": "“血祭以生命换取毁灭性的暴击；劫难则挥出覆盖周身的圆形横扫。”",
        },
        "charred_fire_staff": {
            "category": "法杖",
            "skills": [
                "炬火残光：只要当前手持枯火杖，便持续清除原版负面药水状态。",
                "灵气洪流：长按右键蓄力并移动，松手后留下持续法阵；蓄力最多 2 秒达到最大半径 5 格。",
                "蓄力移动期和持续法阵均每 10 Tick（0.5 秒）造成 250 点伤害，并施加反胃、虚弱 II、失明、缓慢 II、挖掘疲劳 II。",
                "法阵持续 200 Tick（10 秒），技能冷却 600 Tick（30 秒）。",
            ],
            "lore": "“炬火残光｜持有时免疫负面状态。灵气洪流｜长按蓄力，松开后释放持续侵蚀敌人的法阵。”",
        },
        "solar_flare_lance": {
            "category": "近战武器 · 长枪",
            "skills": [
                "耀斑冲锋：长按右键高速突进；真实撞击合法目标时造成玩家当前攻击属性 600% 的魔法伤害。",
                "命中后反弹，并生成残影继续贯穿后方敌人；初始碰撞目标不会被残影重复伤害。",
                "命中施加持续 100 Tick（5 秒）的焦灼状态；随后普通攻击可引爆焦灼，造成范围伤害并引燃目标。",
                "撞上无法跨越的墙体会中断冲锋并反弹。",
            ],
            "lore": "“主动技能：耀斑冲锋。撞击敌人时唤出贯穿残影；再以普通攻击引爆焦灼。”",
        },
        "suspended_clockhand": {
            "category": "近战武器 · 剑",
            "skills": [
                "灵魂汲取：初始灵魂值 10；每次攻击消耗 1 点，击杀目标返还 2 点。",
                "武器附加伤害等于当前灵魂值 ÷4；灵魂降到负数时武器破碎消失。",
                "灵魂上限 2048。满灵魂后长按右键蓄力 100 Tick（5 秒），可秒杀半径 8 格内所有合法目标，并将灵魂重置为 1。",
                "这是唯一允许伤害创造模式玩家的技能，但仍遵守 PVP 与队伍友伤规则。",
            ],
            "lore": "“攻击时汲取灵魂，将灵魂转化为真实的破坏力。对真理的求知，即是与造物主的角力。”",
        },
        "itanis": {
            "category": "远程武器 · 弓",
            "skills": [
                "所有伊塔尼斯箭矢自动追踪；左键切换普通模式与蓄力模式。",
                "普通模式：主箭之外，85% 概率追加 105 点箭、65% 概率追加 180 点箭、25% 概率追加 270 点箭；三项独立判定。",
                "蓄力模式：每 20 Tick（1 秒）生成一支 120 点浮游箭；松手后索敌发射，浮游箭最多保留 100 Tick（5 秒）。",
                "最大蓄力 200 Tick（10 秒）；期间获得抗性提升 II，满蓄力额外发射一支 5000 点贯穿箭。",
                "主箭、额外箭、浮游箭和贯穿箭均使用专用独立结算入口，同一目标被多箭同时命中时每支伤害都能生效。",
            ],
            "lore": "“传说之弓·伊塔尼斯：所有箭矢自动追踪，在疾射与蓄力箭阵之间切换。”",
        },
        "underground_sun": {
            "category": "特殊道具",
            "skills": [
                "长按右键蓄力 40 Tick（2 秒）生成一个环绕光球，最多储存 3 个。",
                "左键发射最早生成的光球；基础攻击 90，爆炸伤害为其 5 倍，即 450 点。",
                "爆炸判定的水平半径为 20 格、垂直半高为 10 格；不破坏方块，并可无视怪物受伤冷却。",
                "发射冷却 5 Tick，光球飞行寿命 100 Tick（5 秒）。",
            ],
            "lore": "“在地底孕育太阳。光球环绕自身提供照明与守护，也可化作不会破坏方块的范围爆炸。”",
        },
        "red_sun": {
            "category": "近战武器 · 剑 / 火球载体",
            "skills": [
                "长按右键最多蓄力 180 Tick（9 秒），生成尺寸从 1 扩大到 12 的烈阳火球。",
                "火球伤害随蓄力从 100 增长到 1250；体积越小初速度越快。",
                "火球触地后沿原方向碾压实体，直至撞墙或动能耗尽才发生大范围爆炸；最大爆炸半径 25 格。",
                "技能冷却 200 Tick（10 秒）。所有攻击附带持续 200 Tick（10 秒）的烧伤：目标切换物品时受到 10 点伤害，且承受伤害提高 20%。",
            ],
            "lore": "“烈阳在蓄力中膨胀，落地后无情碾压；烧伤者每次切换物品都会付出代价。”",
        },
        "corruption_bomb": {
            "category": "投掷道具",
            "skills": [
                "右键投掷；碰撞后影响 5×5×5 的方形范围，对合法目标施加 3 层剧毒。",
                "剧毒稳定期 200 Tick（10 秒）；不直接造成伤害，不破坏方块，也不会影响投掷者。",
                "投射物寿命 200 Tick（10 秒）；物品不可堆叠。",
            ],
            "lore": "“腐败瓶：掷出后，范围内的所有合法目标获得 3 层剧毒印记。”",
        },
    }


def armor_lore(name):
    return f"“{name}｜超新星套装的一部分。星辉随已装备件数逐步苏醒。”"


def build_document():
    data = json.loads(JSON_PATH.read_text(encoding="utf-8"))
    items = data["items"]
    by_id = {item["id"]: item for item in items}
    details = item_details()
    version = read_version()

    armor_ids = [
        "supernova_helmet",
        "supernova_chestplate",
        "supernova_leggings",
        "supernova_boots",
    ]
    for armor_id in armor_ids:
        details[armor_id] = {
            "category": "装备 · 超新星套装",
            "skills": [
                "两件套：每 5 秒使下一次直接近战命中额外追加“当前手持武器攻击力 ×4”的伤害，使用玩家级独立冷却。",
                "四件套：头部两侧出现伴飞星球；攻击命中时发射两枚自动追踪星弹，每枚造成当前武器攻击力 125% 的伤害。星弹可携带重击标志。",
                "当前源码已移除旧版“四件套击杀回满并获得药水增益”逻辑；JSON 中保留的 buffDuration / buffId / buffAmplifier 目前不参与结算。",
            ],
            "lore": armor_lore(by_id[armor_id]["displayName"]["zh_CN"]),
        }

    missing = [item["id"] for item in items if item["id"] not in details]
    if missing:
        raise RuntimeError(f"Missing curated documentation for item ids: {missing}")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Bullet 2"):
        list_style = styles[list_style_name]
        list_style.font.name = FONT_CN
        list_style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        list_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        list_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        list_style.font.size = Pt(10.5)
        list_style.paragraph_format.space_after = Pt(4)
        list_style.paragraph_format.line_spacing = 1.25
        if list_style_name == "List Bullet":
            list_style.paragraph_format.left_indent = Inches(0.375)
            list_style.paragraph_format.first_line_indent = Inches(-0.188)
        else:
            list_style.paragraph_format.left_indent = Inches(0.7)
            list_style.paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(header_paragraph, after=0, line=1)
    header_run = header_paragraph.add_run("DarkGrey · 武器道具装备与印记系统档案")
    set_run_font(header_run, size=8.5, color=MID_GRAY)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(footer_paragraph, after=0, line=1)
    footer_run = footer_paragraph.add_run("第 ")
    set_run_font(footer_run, size=9, color=MID_GRAY)
    add_field(footer_paragraph, "PAGE")
    footer_run2 = footer_paragraph.add_run(" 页")
    set_run_font(footer_run2, size=9, color=MID_GRAY)

    # Editorial cover.
    spacer = doc.add_paragraph()
    style_paragraph(spacer, before=72, after=0)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(kicker, after=18, line=1)
    run = kicker.add_run("DARKGREY · MINECRAFT 1.7.10")
    set_run_font(run, size=10.5, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(title, after=10, line=1.05)
    run = title.add_run("武器、道具、装备\n与印记系统档案")
    set_run_font(run, size=29, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(subtitle, after=42, line=1.2)
    run = subtitle.add_run("当前实现的分类说明、数值口径、技能机制与 Lore")
    set_run_font(run, size=13.5, color=BLUE)

    meta = doc.add_table(rows=2, cols=2)
    meta_values = [
        ("版本快照", f"DarkGrey {version}"),
        ("整理日期", date(2026, 7, 28).isoformat()),
        ("物品数量", f"{len(items)} 件"),
        ("统一印记", "3 种"),
    ]
    for index, (label, value) in enumerate(meta_values):
        cell = meta.rows[index // 2].cells[index % 2]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line=1.15)
        label_run = p.add_run(label + "\n")
        set_run_font(label_run, size=8.5, bold=True, color=MID_GRAY)
        value_run = p.add_run(value)
        set_run_font(value_run, size=12, bold=True, color=NAVY)
    set_table_geometry(meta, [4680, 4680], indent_dxa=120)
    set_table_borders(meta, color="D7E1EC", size="8")

    cover_note = doc.add_paragraph()
    cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(cover_note, before=54, after=0, line=1.2)
    run = cover_note.add_run("以 Jar 内置 RPG 注册表与当前 Java 实现为准")
    set_run_font(run, size=9.5, italic=True, color=MID_GRAY)

    doc.add_page_break()

    doc.add_heading("1. 阅读说明与系统总览", level=1)
    add_note_box(
        doc,
        "数据口径",
        "“基础攻击力”取自当前发布注册表的 damage 字段；技能伤害单独写入“特殊技能”。装备与弹药的攻击力标为“不适用”或 0，不将套装伤害误写成物品基础攻击力。",
    )
    add_label_paragraph(
        doc,
        "文档范围：",
        f"当前 Jar 发布注册表中的 {len(items)} 件物品、3 种统一印记，以及与武器直接相关的专属状态。",
    )
    add_label_paragraph(
        doc,
        "Tick 换算：",
        "Minecraft 默认 20 Tick = 1 秒。本文同时给出 Tick 与秒数，便于核对配置和玩法。",
    )
    add_label_paragraph(
        doc,
        "Lore 口径：",
        "优先整理当前组件实际加入物品提示框的文本；没有独立提示的物品会明确注明，并用其当前用途作简短介绍。",
    )

    doc.add_heading("分类导航", level=2)
    categories = [
        ("近战、远程与法杖武器", "12 件"),
        ("投掷与特殊道具", "3 件"),
        ("装备与套装", "4 件"),
        ("弹药", "1 件"),
        ("统一印记系统", "剧毒、骨折、碎骨"),
        ("武器专属状态", "焦灼、烧伤"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "分类"
    table.rows[0].cells[1].text = "内容"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=WHITE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for label, value in categories:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=9.5, color=DARK)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [5600, 3760], indent_dxa=120)
    set_table_borders(table)

    doc.add_heading("全物品速查", level=2)
    overview = doc.add_table(rows=1, cols=5)
    headers = ["名称", "分类", "基础攻击", "耐久", "核心机制"]
    for index, header in enumerate(headers):
        overview.rows[0].cells[index].text = header
        set_cell_shading(overview.rows[0].cells[index], BLUE)
        for run in overview.rows[0].cells[index].paragraphs[0].runs:
            set_run_font(run, size=8.8, bold=True, color=WHITE)
    summary_text = {
        "bone_flask": "投掷骨刺场 / 骨折",
        "bone_crusher": "三击叠加骨折",
        "erebus": "球形剧毒领域",
        "cleave_sword": "5 秒重击",
        "vampire_blade": "15% 吸血 + 重击",
        "supernova_helmet": "超新星套装",
        "supernova_chestplate": "超新星套装",
        "supernova_leggings": "超新星套装",
        "supernova_boots": "超新星套装",
        "rainbow_bow": "三段蓄力",
        "law_of_cycles": "30 支光矢箭阵",
        "arrow_of_light": "圆环之理弹药",
        "calamity_scythe": "血祭 + 圆形横扫",
        "charred_fire_staff": "免疫负面 + 法阵",
        "solar_flare_lance": "冲锋 / 焦灼引爆",
        "suspended_clockhand": "灵魂成长 / 范围秒杀",
        "itanis": "追踪多箭 / 浮游箭阵",
        "underground_sun": "储存与发射光球",
        "red_sun": "蓄力碾压火球",
        "corruption_bomb": "范围施加剧毒",
    }
    for item in items:
        row = overview.add_row()
        item_id = item["id"]
        attack = "不适用" if item_id in armor_ids or item_id == "arrow_of_light" else str(item["damage"])
        values = [
            item["displayName"]["zh_CN"],
            details[item_id]["category"],
            attack,
            str(item.get("durability", 0)),
            summary_text[item_id],
        ]
        for index, value in enumerate(values):
            row.cells[index].text = value
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in row.cells[index].paragraphs[0].runs:
                set_run_font(run, size=8.4, color=DARK)
            style_paragraph(row.cells[index].paragraphs[0], after=0, line=1.1)
        if len(overview.rows) % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "F8FAFC")
    set_repeat_table_header(overview.rows[0])
    set_table_geometry(overview, [1900, 2250, 1100, 900, 3210], indent_dxa=120)
    set_table_borders(overview, color="D4DAE2", size="5")

    weapon_order = [
        "bone_crusher",
        "cleave_sword",
        "vampire_blade",
        "calamity_scythe",
        "suspended_clockhand",
        "red_sun",
        "rainbow_bow",
        "law_of_cycles",
        "itanis",
        "erebus",
        "charred_fire_staff",
        "solar_flare_lance",
    ]
    utility_order = ["bone_flask", "corruption_bomb", "underground_sun"]

    def add_item_entry(item_id, heading_level=2, force_new_page=False):
        item = by_id[item_id]
        detail = details[item_id]
        heading = doc.add_heading(item["displayName"]["zh_CN"], level=heading_level)
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.page_break_before = force_new_page
        attack_text = (
            "不适用（装备）"
            if item_id in armor_ids
            else "不适用（弹药）"
            if item_id == "arrow_of_light"
            else str(item["damage"])
        )
        max_stack = item.get("maxStackSize", 1 if item_id in ("bone_flask", "corruption_bomb") else "默认")
        add_meta_table(
            doc,
            [
                ("内部 ID", item_id),
                ("分类 / 类型", f"{detail['category']} / {item['type']}"),
                ("基础攻击力", attack_text),
                ("耐久度配置", item.get("durability", 0)),
                ("附魔", enchantments_text(item)),
                ("最大堆叠", max_stack),
            ],
        )
        label = doc.add_paragraph()
        style_paragraph(label, before=8, after=4, line=1.1, keep_with_next=True)
        run = label.add_run("特殊技能")
        set_run_font(run, size=10.5, bold=True, color=GREEN)
        for skill in detail["skills"]:
            add_bullet(doc, skill)
        add_label_paragraph(doc, "Lore（游戏内提示整理）：", detail["lore"], color=PURPLE)

    doc.add_page_break()
    doc.add_heading("2. 武器", level=1)
    add_note_box(
        doc,
        "阅读提示",
        "武器条目先列注册表基础攻击力，再列技能伤害。弓类技能的箭矢伤害、法杖的持续伤害以及主动技能倍率均不会并入基础攻击力。",
    )
    standalone_weapon_pages = {
        "red_sun",
        "law_of_cycles",
        "itanis",
        "erebus",
        "charred_fire_staff",
        "solar_flare_lance",
    }
    for item_id in weapon_order:
        add_item_entry(item_id, force_new_page=item_id in standalone_weapon_pages)

    doc.add_page_break()
    doc.add_heading("3. 投掷与特殊道具", level=1)
    for item_id in utility_order:
        add_item_entry(item_id, force_new_page=item_id in {"corruption_bomb", "underground_sun"})

    doc.add_page_break()
    doc.add_heading("4. 装备与超新星套装", level=1)
    add_note_box(
        doc,
        "当前实现",
        "四件超新星装备共用同一套装组件。当前自定义护甲材质的基础护甲点数组为 0/0/0/0，因此本章重点记录套装技能，不将耐久度误认为防御力。",
        fill="FFF8E8",
        accent=GOLD,
    )
    armor_table = doc.add_table(rows=1, cols=5)
    armor_headers = ["名称", "内部 ID", "类型", "基础攻击", "耐久"]
    for index, header_text in enumerate(armor_headers):
        armor_table.rows[0].cells[index].text = header_text
        set_cell_shading(armor_table.rows[0].cells[index], BLUE)
        for run in armor_table.rows[0].cells[index].paragraphs[0].runs:
            set_run_font(run, size=9, bold=True, color=WHITE)
    for armor_id in armor_ids:
        item = by_id[armor_id]
        cells = armor_table.add_row().cells
        values = [
            item["displayName"]["zh_CN"],
            armor_id,
            item["type"],
            "不适用",
            str(item["durability"]),
        ]
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[index].paragraphs[0].runs:
                set_run_font(run, size=9, color=DARK)
            style_paragraph(cells[index].paragraphs[0], after=0, line=1.1)
    set_repeat_table_header(armor_table.rows[0])
    set_table_geometry(armor_table, [1750, 2950, 1200, 1500, 1960], indent_dxa=120)
    set_table_borders(armor_table)

    shared_label = doc.add_paragraph()
    style_paragraph(shared_label, before=12, after=4, line=1.1, keep_with_next=True)
    shared_run = shared_label.add_run("共享特殊技能")
    set_run_font(shared_run, size=10.5, bold=True, color=GREEN)
    for skill in details["supernova_helmet"]["skills"]:
        add_bullet(doc, skill)
    doc.add_heading("逐件 Lore", level=2)
    for armor_id in armor_ids:
        item_name = by_id[armor_id]["displayName"]["zh_CN"]
        add_bullet(doc, f"{item_name}：{details[armor_id]['lore']}")

    doc.add_page_break()
    doc.add_heading("5. 弹药与辅助物品", level=1)
    add_item_entry("arrow_of_light")

    doc.add_page_break()
    doc.add_heading("6. 统一印记系统", level=1)
    add_note_box(
        doc,
        "框架规则",
        "衰减模式、衰减间隔和每次衰减层数是印记固有属性；稳定期由每一次施加行为显式传入。调用方未指定新数值时，也应显式传入该印记默认稳定期。",
    )
    doc.add_heading("6.1 生命周期与衰减模式", level=2)
    add_bullet(doc, "稳定期：印记被施加或刷新后进入稳定倒计时；不同武器可以为同一种印记传入不同稳定时间。")
    add_bullet(doc, "持续型衰减：稳定期结束后进入一段衰减期，按照固定间隔持续减少层数，直到归零或被重新施加。")
    add_bullet(doc, "瞬时型衰减：稳定期结束时立即减少固定层数；若仍有层数，则按该实例稳定期重新开始倒计时。")
    add_bullet(doc, "重新施加：覆盖该印记实例的稳定期参数、重置稳定倒计时，并退出当前衰减状态。")

    doc.add_heading("6.2 剧毒", level=2)
    add_meta_table(
        doc,
        [
            ("内部 ID", "poison"),
            ("类型", "持续型（CONTINUOUS）"),
            ("最大层数", "99"),
            ("默认稳定期", "200 Tick（10 秒）"),
            ("衰减规则", "稳定期结束后，每 10 Tick（0.5 秒）减少 10 层"),
            ("周期触发", "每 10 Tick（0.5 秒）触发一次，伤害等于当前剧毒层数"),
        ],
    )
    add_label_paragraph(doc, "作用：", "剧毒在施加时立即按当前层数造成一次伤害，随后每 0.5 秒再次按当前层数造成伤害；当前配置允许其无视受伤无敌帧。")
    add_label_paragraph(doc, "Lore：", "“受到剧毒的侵蚀。”")
    add_label_paragraph(doc, "主要来源：", "厄瑞波斯、腐败瓶。")

    doc.add_heading("6.3 骨折", level=2)
    add_meta_table(
        doc,
        [
            ("内部 ID", "fracture"),
            ("类型", "瞬时型（INSTANT）"),
            ("最大层数", "5"),
            ("默认稳定期", "100 Tick（5 秒）"),
            ("衰减规则", "稳定期结束时减少 1 层；若仍有层数，则重新倒计时"),
            ("属性影响", "每层降低 10% 移动速度，最高按 95% 减速上限保护"),
        ],
    )
    add_label_paragraph(doc, "作用：", "层数越高移动越慢；达到 5 层时自动维持“碎骨”，离开满层后停止该满层维持来源。")
    add_label_paragraph(doc, "Lore：", "“每层降低移动速度；达到 5 层时诱发严重碎骨，导致移动受损。”")
    add_label_paragraph(doc, "主要来源：", "粉碎之骨、碎骨瓶骨刺场。")

    doc.add_heading("6.4 碎骨", level=2)
    add_meta_table(
        doc,
        [
            ("内部 ID", "shattered_bone"),
            ("类型", "瞬时型（INSTANT），最大 1 层"),
            ("默认稳定期", "60 Tick（3 秒）"),
            ("独立持续", "独立施加时按本次稳定期计时"),
            ("骨折维持", "骨折满 5 层时可持续维持；骨折离开满层后解除该维持来源"),
            ("移动惩罚", "移动达到阈值时每 5 Tick 最多受到 5 点伤害"),
        ],
    )
    add_label_paragraph(
        doc,
        "作用：",
        "携带者移动会承受伤害；受到物理攻击时，骨骼碎片向前方扇形溅射，造成攻击伤害 225% 的穿甲伤害，并向附近合法目标扩散持续 60 Tick（3 秒）的碎骨。",
    )
    add_label_paragraph(
        doc,
        "Lore：",
        "“由严重骨折引发。移动会承受巨大伤害；受到物理攻击时，骨骼碎片向外溅射并扩散碎骨状态。”",
    )
    add_label_paragraph(doc, "主要来源：", "骨折达到 5 层；碎骨受击后的碎片扩散。")

    doc.add_heading("6.5 印记显示系统", level=2)
    add_bullet(doc, "自身 HUD：每种印记使用独立 120×32 底板；每列最多 4 个，第 5 个自动换列。")
    add_bullet(doc, "实体头顶：每排最多 3 个印记，按新印记从右侧加入并形成居中的多行阵列。")
    add_bullet(doc, "目标面板：准星指向实体时，每种印记只占一行，显示名称、层数、稳定时间与预计伤害等摘要。")
    add_bullet(doc, "衰减警报：进入衰减状态后，图标与相关数字同步闪烁，稳定倒计时显示为 0:00。")

    doc.add_page_break()
    doc.add_heading("7. 武器专属状态（非统一印记）", level=1)
    add_note_box(
        doc,
        "系统边界",
        "下列状态在玩法文字中可能被称作“印记”或“异常”，但它们不在 MarkRegistry 的统一印记类型中，不能与剧毒、骨折、碎骨使用同一套稳定期/衰减接口。",
        fill="FFF4F2",
        accent=RED,
    )
    doc.add_heading("7.1 焦灼印记", level=2)
    add_label_paragraph(doc, "来源：", "耀斑长枪的冲锋撞击与残影命中。")
    add_label_paragraph(doc, "持续时间：", "100 Tick（5 秒）。")
    add_label_paragraph(doc, "作用：", "普通攻击命中带焦灼的敌人时引爆，造成范围伤害并引燃目标。焦灼由专属 ScorchedMarkTracker 管理。")
    add_label_paragraph(doc, "Lore：", "“再次普通攻击可引爆印记，造成周围范围性的毁灭伤害，并引燃目标。”")

    doc.add_heading("7.2 烧伤异常", level=2)
    add_label_paragraph(doc, "来源：", "红日的所有攻击。")
    add_label_paragraph(doc, "持续时间：", "200 Tick（10 秒）。")
    add_label_paragraph(doc, "作用：", "携带者每次切换物品时受到 10 点伤害；同时其承受伤害提高 20%。")
    add_label_paragraph(doc, "Lore：", "“烧伤者每次切换物品都会受伤，并失去一部分防御能力。”")

    doc.add_page_break()
    doc.add_heading("8. 维护备注与数据来源", level=1)
    add_bullet(doc, "物品名称、类型、基础攻击力、耐久、附魔和组件参数：Jar 内置 rpg_items.json。")
    add_bullet(doc, "特殊技能与 Lore：当前 Java 组件实现、组件 Tooltip 与 zh_CN.lang。")
    add_bullet(doc, "统一印记：MarkRegistry 注册项、PoisonMarkType、FractureMarkType、ShatteredBoneMarkType 与 Config 默认值。")
    add_bullet(doc, "文档记录的是当前源码快照，不代表以后热更新或新版本仍保持相同数值。")
    add_note_box(
        doc,
        "已发现的说明差异",
        "超新星 JSON/Excel 仍保留旧击杀增益参数，但当前源码已移除旧四件套击杀逻辑；本文已按实际运行代码记录为伴飞星弹。骨折中文描述沿用旧措辞，本文的稳定与衰减规则按当前 MarkType 实现整理。",
        fill="FFF8E8",
        accent=GOLD,
    )

    # Core document metadata and update field instruction.
    doc.core_properties.title = "DarkGrey武器道具装备与印记系统说明"
    doc.core_properties.subject = f"DarkGrey {version} 当前系统档案"
    doc.core_properties.author = "DarkGrey 项目"
    doc.core_properties.keywords = "DarkGrey,Minecraft 1.7.10,武器,道具,装备,印记"

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"CREATED={OUTPUT_PATH}")
    print(f"ITEMS={len(items)}")
    print(f"VERSION={version}")


if __name__ == "__main__":
    build_document()
